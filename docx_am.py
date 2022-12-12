import nltk
from nltk.tokenize import sent_tokenize
import srt

nltk.download ('punkt')


def getSrt (srt_fn) :
    NUM_PARAGRAPHS = 2

    with open (srt_fn) as rf :
        buf = rf.read ()    
        srt_data = srt.parse (buf)
        srt_list = list (srt_data)

    last_time = srt_list [-1].end.total_seconds ()
    split_length = int (last_time / NUM_PARAGRAPHS + 2)

    print (last_time)
    print ('half_time in seconds :', split_length)
    
    return srt_list, last_time, split_length
###


### 시간 기준으로 전체 문서를 나눔
def getParts (srt_list, split_length) : ### length of a part in seconds
    split_time = split_length
    text_lines = []
    start_time = -1
    
    for sub in srt_list :
        index   = sub.index
        content = sub.content
        st      = sub.start.total_seconds ()
        et      = sub.end.total_seconds ()
        
        if start_time == -1 :
            start_time = st

        if st < split_time :
            if content != 'CHEERING AND APPLAUSE' :
                text_lines.append (content)
            continue
        ###
        
        #print ('###', index, st, content)
        text = ' '.join (text_lines)
        yield text, start_time
        
        text_lines = []
        split_time += split_length
        start_time = -1
    ###
    
    text = ' '.join (text_lines)
    yield text, start_time
###


def getPartsText (srt_list, split_length) :
    """ parts : list of split text. one part for one trascriber
        parts_start : start time of the split text
    """
    
    seg_length = 600 ### 10 minutes
    parts = []
    parts_start = []

    it_parts = getParts (srt_list, split_length)

    for text, start_time in it_parts :
        if text == '' : break

        parts.append (text)
        parts_start.append (start_time)

        print ('words :', len (text.split (' ')))
        print ('characters :', len (text))

        sentences = sent_tokenize (text)
        #print (len (sentences))
        #print ('\n'.join (sentences))
        print ('='*80)
    ###
    
    return parts, parts_start
###


def splitToParagraphs (text, split_length=1000) :
    """ split text into small paragraphs : around 10000 characters each
        just to make it easy to read the text
    """
    
    tlen = len (text) ### total number of characters
    mods = int (tlen / split_length + 0.5) ### number of paragraphs
    para_size = int (tlen / mods) + 1 ### characters per paragraph
    
    sentences = sent_tokenize (text)
    
    paragraphs = [] ### result
    total_len = 0
    buf = []
    
    for sentence in sentences :
        total_len += len (sentence)
        buf.append (sentence)
        
        if total_len >= para_size :
            paragraph = ' '.join (buf)
            paragraphs.append (paragraph)
            
            total_len = 0
            buf = []
        ###
    ###

    if len (buf) > 0 :
        paragraph = ' '.join (buf)
        paragraphs.append (paragraph)

    return paragraphs
###
        

def getDocText (parts) :
    """ doc_text : [ part1, part2, ... ] = [ [ para1, para2, ...], [ para3, para4, ...], ... ]
    """
    
    doc_text = []

    for text in parts :
        paragraphs = splitToParagraphs (text)
        doc_text.append (paragraphs)
    ###
    
    return doc_text
###


from docx import Document
from docx.shared import Pt
import time

# save to docx file
@staticmethod
def saveToDocx (title, doc_text, parts_start, time_length, filename) :
    document = Document()
    para = document.add_paragraph ()

    ### title
    run = para.add_run (title)
    run.font.size = Pt (16) # font size
    run.font.name = 'Arial' # font name
    run.bold = True

    ### body
    for idx, (paragraphs, start_time) in enumerate (zip (doc_text, parts_start)) :
        if idx == 0 :
            time_string = 'Length ' + time.strftime ('%H:%M:%S', time.gmtime (time_length))
        else :
            time_string = time.strftime ('%H:%M:%S', time.gmtime (start_time))
        
        body = document.add_paragraph ()
        run = body.add_run ('\n')
        run = body.add_run (time_string)
            
        run.font.size = Pt (14) # font size
        run.font.name = 'Arial' # font name
        run.bold = True

        for para_text in paragraphs :
            body = document.add_paragraph ()
            run = body.add_run ('\n')

            #sentences = text.split ('\n')
            run = body.add_run (para_text)
            run.font.size = Pt (14) # font size
            run.font.name = 'Arial' # font name

        run = body.add_run ('\n')
    ###
    
    ### save document
    document.save (filename)
###


def makeDocx (srt_file, doc_file) :
    """ srt_file : input srt file
        doc_file : result docx file
    """
    
    srt_list, last_time, split_length = getSrt (srt_file)
    parts, parts_start = getPartsText (srt_list, split_length)
    doc_text = getDocText (parts)

    title = '.'.join (doc_file.split ('/') [-1].split ('.') [:-1])
    saveToDocx (title, doc_text, parts_start, last_time, doc_file)
###


if __name__ == '__main__' :
    srt_file = 'data/sample-0.srt'
    srt_file = 'data/dream/dream.srt'
    doc_file = 'data/dream/dream.docx'

    #doc_file = output_dir + name + '.docx'
    makeDocx (srt_file, doc_file)


