#from google.colab import drive

from docx import Document
from docx.shared import Pt
from docx_am import makeDocx

import ffmpeg
import glob
import numpy as np
import os
from os.path import basename
import re
import time
import torch
from whisper_utils import write_txt, write_srt
from zipfile import ZipFile
# from br2us import british_to_american

# FASTER = True ### use faster-whisper
# print ('FASTER :' , FASTER, 'whisper_am.py')

import whisper
from faster_whisper import WhisperModel


try:
    from google.colab import drive
    from google.colab import files
    from google.colab import runtime

    is_google_colab = True

except ImportError:
    is_google_colab = False


def mountDrive () :
    if is_google_colab == False :
        print ('Not a Google Colab environment !!!')
        return False

    google_drive = '/content/drive/MyDrive/'
    isExist = os.path.exists (google_drive)
    if not isExist :
        drive.mount('/content/drive/')
        print ('Google Driver mounted.')
    
    return True
###

            
class WhisperAM :
    def __init__ (self, model_name, data_dir, initial_prompt="", doc_download=True, beam_size=5, FASTER=True) :
        self.model_name     = model_name
        self.data_dir       = data_dir
        self.out_dir        = data_dir
        self.initial_prompt = initial_prompt
        self.doc_download   = doc_download
        self.beam_size      = beam_size
        self.FASTER         = FASTER
        
        isExist = os.path.exists (self.out_dir)
        if not isExist :
            os.makedirs (self.out_dir, exist_ok=False)
    ###


    ### https://betterprogramming.pub/openais-whisper-tutorial-42140dd696ee
    def checkCuda (self) :
        # make sure if it shows 'cuda' which means gpu ready
        # if not, change the runtime type : menu > runtime > change runtime type > select GPU, then start from the beginning again

        torch.cuda.is_available ()
        self.DEVICE = "cuda" if torch.cuda.is_available () else "cpu"
            
        return self.DEVICE
    ###


    #@title Whisper model selection : medium.en (default)
    def loadModel (self) :
        if self.FASTER :
            self.model = WhisperModel (self.model_name, device=self.DEVICE, compute_type="float16") ### faster whisper
        else :
            self.model = whisper.load_model (self.model_name) ### default whisper

        #self.model = whisper.load_model (self.model_name, device=self.DEVICE)
    ###


    def getText (self, segments) :
        text = []
        for segment in segments:
            text.append (segment['text'].strip())

        return '\n'.join (text)
    ###


    #@title transcribe function definition
    def transcribe (self, mp3_file) :
        options = dict (language='English', beam_size=self.beam_size, best_of=5)
        transcribe_options = dict (task="transcribe", **options)
        
        if self.initial_prompt != "" :
            self.result = self.model.transcribe (mp3_file, verbose=False, initial_prompt=self.initial_prompt, **transcribe_options)
        else :
            self.result = self.model.transcribe (mp3_file, verbose=False, **transcribe_options)

        return self.result
    ###        
    

    #@title faster transcribe function definition
    def transcribeFaster (self, mp3_file) :
        options = dict (language='English', beam_size=self.beam_size, best_of=5)
        transcribe_options = dict (task="transcribe", **options)
        
        if self.initial_prompt != "" :
            segments, info = self.model.transcribe (mp3_file, initial_prompt=self.initial_prompt, **transcribe_options)
        else :
            segments, info = self.model.transcribe (mp3_file, **transcribe_options)

        segments = list (segments)
        segments2 = []
        text = []

        for s in segments :
            seg = { 'start': s.start, 'end': s.end, 'text': s.text.strip () }
            segments2.append (seg)
            text.append (s.text.strip ())

        self.result = dict (segments=segments2, text=' '.join (text))

        return self.result
    ###


    def saveResult (self, name, output_dir, result) :
        # text = self.result ['text']
        segments = self.result ['segments']

        ### save temp SRT : caption
        srt_file_tmp = output_dir + name + '.srt.tmp'
        with open (srt_file_tmp, "w", encoding="utf-8") as wf:
            write_srt (segments, file=wf)

        ### postprocess srt
        with open (srt_file_tmp) as rf:
            text = rf.read ()
            text2 = self.postprocess (text)

        ### save SRT : caption
        srt_file = output_dir + name + '.srt'
        with open (srt_file, "w", encoding="utf-8") as wf:
            wf.write (text2)
            os.remove (srt_file_tmp)
            
        ### post processing : apply macros
        text = self.getText (segments)
        text2 = self.postprocess (text)

        txt_file = output_dir + name + '.txt'
        with open (txt_file, 'w') as wf :
            wf.write (text2)
            #os.remove (tmp_file)


        doc_file = output_dir + name + '.docx'
        #self.saveToDocx (text2, doc_file)
        makeDocx (srt_file, doc_file)

        return result    
    ###



    #@title postprocessing : macro, save to docx


    ### macro definition
    rdict = { '1000s': 'thousands', 'a collective consciousness': 'the collective consciousness', 'Ascended Master': 'ascended master', 'a golden age': 'the golden age', 'And so ': '', 'And so, ': '', 'and so forth and so on': 'and so forth', "aren't": 'are not', 'behaviour': 'behavior', "can't": 'cannot', 'Christ to it': 'Christhood', 'colour': 'color', 'conscious you': 'Conscious You', "couldn't": 'could not', 'defence': 'defense', "didn't": 'did not', "doesn't": 'does not', "don't": 'do not', 'Earth': 'earth', 'fall and beaks': 'fallen beings', 'flavour': 'flavor', 'freewill': 'free will', 'Golden Age': 'golden age', "haven't": 'have not', "he's": 'he is', "I'm": 'I am', "isn't": 'is not', "it's": 'it is', "It's": 'It is', 'karmic board': 'Karmic Board', 'labour': 'labor', 'offence': 'offense', 'out picture': 'out-picture', 'outpicture': 'out-picture', 'realise': 'realize', 'saviour': 'savior', 'self awareness': 'self-awareness', "shouldn't": 'should not', 'summit lighthouse': 'Summit Lighthouse', "that's": 'that is', "That's": 'That is', "there's": 'there is', "There's": 'There is', "they're": 'they are', "wasn't": 'was not', "we're": 'we are', "We're": 'We are', "weren't": 'were not', "we've": 'we have', "What's": 'What is', "what's": 'what is', "wouldn't": 'would not', "won't": 'will not', "you'll": 'you will', "you're": 'you are', "You're": 'You are', "here's": "here is", "Here's": "Here is", "where's": "where is", "who's": "who is", "Who's": "Who is", "I've": "I have", 'So ': '', 'So, ': '', 'Well ': 'Well, ', 'separate cell': 'separate self', 'separate cells': 'separate selves', 'followers bodies': "four lower bodies", "matter light": "Ma-ter light", "log in": "lock in", "divine plan": "Divine Plan", "other ascended": "unascended", "an embodiment": "in embodiment" }


    @staticmethod
    def multiple_replace (adict, text):
        ### Create a regular expression from all of the dictionary keys
        regex = re.compile ( "|".join (map (re.escape, adict.keys ( ))) )

        ### For each match, look up the corresponding value in the dictionary
        return regex.sub (lambda match: adict [match.group (0)], text)
    ###


    def postprocess (self, text) :
        ### macro processing
        callback = lambda pat: pat.group(2).upper()
        text2 = re.sub ('(So|So,|And so|And so,) ([a-zA-Z])', callback, text) ### capitalize after So, And so

        # text2 = british_to_american (text2) ### convert British to American English
        text2 = self.multiple_replace (self.rdict, text2)

        return text2
    ###


    ### save to docx file
    ### not used anymore. now using makeDocx in docx_am.py
    @staticmethod
    def saveToDocx (text, filename) :
        document = Document()
        para = document.add_paragraph ()

        ### title
        title = '.'.join (filename.split ('/') [-1].split ('.') [:-1])
        run = para.add_run (title)
        run.font.size = Pt (16) # font size
        run.font.name = 'Arial' # font name
        run.bold = True

        ### body
        body = document.add_paragraph ()
        run = body.add_run ('\n')

        sentences = text.split ('\n')
        run = body.add_run (' '.join (sentences))
        run.font.size = Pt (14) # font size
        run.font.name = 'Arial' # font name

        ### save document
        document.save (filename)
    ###        


    ### get input audio data list in 'data' folder
    @staticmethod
    def getFiles (data_dir) :
        files1 = glob.glob (data_dir + '/*.mp3')
        files2 = glob.glob (data_dir + '/*.m4a')

        input_data = files1 + files2
        return input_data
    ###
    
    
    ### down sample input audio to 16k
    ### https://www.programcreek.com/python/example/117479/ffmpeg.output
    @staticmethod
    def downSample (in_name, out_name) :
        try :
            audio  = ffmpeg.input  (in_name).audio
            stream = ffmpeg.output (audio, out_name, **{'ar': '16000', 'acodec': 'mp3'}).overwrite_output ()
            out    = ffmpeg.run    (stream, capture_stdout=True, capture_stderr=True)
        except Exception as e :
            print (e)
    ###
    
    
    #@title unassign google colab resource
    @staticmethod
    def unassignColab () :
        ### shutdown colab runtime. works well
        if is_google_colab :
            print ('Google Colab resource unassigned.')
            runtime.unassign ()
    ###
    

    ### download docx files
    @staticmethod
    def downloadFiles (data_path) :
        if is_google_colab :
            docx_files = glob.glob (data_path + '*.docx')

            for doc_file in docx_files :
                files.download (doc_file)
    ###


    ### download docx files in a zip file : docx.zip
    @staticmethod
    def downloadFilesZip (data_path) :
        if is_google_colab :
            docx_files = glob.glob (data_path + '*.docx')

            with ZipFile ('docx.zip', 'w') as zip_obj :
                for doc_file in docx_files :
                    zip_obj.write (doc_file, basename (doc_file))
            ###

            files.download ('docx.zip')
    ###


    ### main program
    def run (self) :
        self.loadModel ()
        input_data = self.getFiles (self.data_dir)
        
        for source_file_name in input_data :
            print (source_file_name)

            target_name = '.'.join (source_file_name.split ('/') [-1].split ('.') [:-1])
            if is_google_colab :
                mp3_file_16k = '/content/' + target_name + '-16k.mp3'
            else :
                mp3_file_16k = self.data_dir + '/' + target_name + '-16k.mp3'

            ### don't need to downsample. original file is better for whisper
            self.downSample (source_file_name, mp3_file_16k)

            ### trascribing
            if self.FASTER :
                result = self.transcribeFaster (mp3_file_16k)
            else :
                result = self.transcribe (mp3_file_16k)

            save_result = self.saveResult (target_name, self.data_dir, result)
            text = save_result ['text']
            print (text [:50], '...', text [-50:])
            print ()
            
            os.remove (mp3_file_16k)
        ### for

        if is_google_colab and self.doc_download :
            self.downloadFilesZip (self.data_dir)
    ### main
### class




def test () :
    #@title run Whisper

    data_dir = 'data/'

    ds = mountDrive ()
    google_drive = '/content/drive/MyDrive/' if ds else './'

    data_path   = google_drive + data_dir
    print (data_path)

    model_name   = 'tiny.en'
    doc_download = False
    
    am = WhisperAM (model_name, data_path, doc_download=doc_download)
    am.checkCuda ()

    print ('Start :', time.strftime('%X %x %Z'))
    am.run ()
    print ('End  :', time.strftime('%X %x %Z'))
###


if __name__ == '__main__' :
    test ()
